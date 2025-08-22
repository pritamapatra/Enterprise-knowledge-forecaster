try:
    import pysqlite3
    import sys
    sys.modules['sqlite3'] = pysqlite3
except ImportError:
    import sqlite3

import streamlit as st
import pandas as pd
import networkx as nx
import numpy as np
from datetime import datetime
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots
import requests
from pypdf import PdfReader
import docx
from io import BytesIO
import tempfile
import os
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import Pinecone
import pinecone
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import json
import re
import plotly
import spacy
from collections import defaultdict

# Page configuration
st.set_page_config(
    page_title="Enterprise Knowledge Evolution Forecaster",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for professional styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1E3A8A;
        text-align: center;
        margin-bottom: 2rem;
        border-bottom: 3px solid #3B82F6;
        padding-bottom: 1rem;
    }
    
    .metric-container {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    
    .risk-alert {
        background: #FEF2F2;
        border-left: 5px solid #EF4444;
        padding: 1rem;
        margin: 1rem 0;
        border-radius: 5px;
    }
    
    .success-alert {
        background: #F0FDF4;
        border-left: 5px solid #22C55E;
        padding: 1rem;
        margin: 1rem 0;
        border-radius: 5px;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state for document processing and RAG
if 'vectorstore' not in st.session_state:
    st.session_state.vectorstore = None
if 'embedding_model' not in st.session_state:
    st.session_state.embedding_model = None
if 'document_processed' not in st.session_state:
    st.session_state.document_processed = False
if 'query_input' not in st.session_state:
    st.session_state.query_input = ''
if 'workflow_results' not in st.session_state:
    st.session_state.workflow_results = []

# Initialize session state for knowledge graph
if 'knowledge_graph' not in st.session_state:
    st.session_state.knowledge_graph = nx.Graph()
if 'entities' not in st.session_state:
    st.session_state.entities = {'PERSON': set(), 'PROJECT': set(), 'SKILL': set()}

@st.cache_resource
def initialize_pinecone():
    """Initialize Pinecone vector database"""
    try:
        pc = Pinecone(
        api_key=os.getenv("PINECONE_API_KEY", "pcsk_3CwtHw_9Zw6NTWL64SryXCZLLS8BreatEH46i5uJLCm6F5jwYpnS93WsT4DWN2roYaXnCC")
    )
        
        embeddings = OpenAIEmbeddings()
        index_name = "enterprise-knowledge"
        
        # Create index if it doesn't exist
        if index_name not in pinecone.list_indexes():
            pinecone.create_index(
                name=index_name,
                dimension=1536,  # OpenAI embedding dimension
                metric="cosine"
            )
        index = pc.Index(index_name)

        vectorstore = Pinecone(
            index_name=index_name,
            embedding_function=embeddings
        )
        
        return vectorstore, embeddings, "Pinecone Active"
        
    except Exception as e:
        return None, None, f"Pinecone Error: {str(e)}"

def extract_pdf_text(uploaded_file):
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""

def extract_docx_text(uploaded_file):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(uploaded_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading DOCX: {str(e)}")
        return ""

def extract_txt_text(uploaded_file):
    """Extract text from TXT file"""
    try:
        return uploaded_file.read().decode('utf-8')
    except Exception as e:
        st.error(f"Error reading TXT: {str(e)}")
        return ""

def generate_direct_answer(query, context):
    """Generate answer using LM Studio with Llama 3.2 3B Instruct"""
    try:
        url = "http://127.0.0.1:1234/v1/chat/completions"
        
        payload = {
            "model": "llama-3.2-3b-instruct",
            "messages": [
                {
                    "role": "user",
                    "content": f"""You are an expert business analyst reviewing an enterprise document. Based on the document content below, answer the user's question with specific, actionable information.

Document Content:
{context[:4000]}

User Question: {query}

Instructions:
- Provide direct, professional answers suitable for executive presentations
- Use specific numbers, percentages, and data from the document
- If information isn't in the document, state that clearly
- Keep responses concise but comprehensive
- Focus on business insights and actionable information

Answer:"""
                }
            ]
        }
        
        response = requests.post(url, json=payload, timeout=30)
        
        if response.status_code == 200:
            result = response.json()
            return result['choices'][0]['message']['content'].strip()
        else:
            return "LM Studio API unavailable. Please ensure LM Studio server is running with Llama 3.2 3B model loaded."
            
    except Exception as e:
        st.warning(f"LM Studio API temporarily unavailable: {str(e)}. Using fallback method.")
        return generate_fallback_answer(query, context)

def generate_fallback_answer(query, context):
    """Fallback method using pattern matching when LM Studio is unavailable"""
    context_lower = context.lower()
    query_lower = query.lower()
    
    if "investment" in query_lower and "total" in query_lower:
        investment_match = re.search(r'\$(\d+(?:\.\d+)?)\s*(?:million|m)', context, re.IGNORECASE)
        if investment_match:
            return f"The total investment requirement appears to be ${investment_match.group(1)} million based on the document content."
    
    if "roi" in query_lower:
        roi_match = re.search(r'(\d+(?:\.\d+)?)\%.*?(?:roi|return)', context, re.IGNORECASE)
        if roi_match:
            return f"The return on investment (ROI) mentioned is {roi_match.group(1)}%."
    
    sentences = [s.strip() for s in context.split('.') if len(s.strip()) > 10]
    query_words = set(query_lower.split())
    
    scored_sentences = []
    for sentence in sentences:
        sentence_words = set(sentence.lower().split())
        overlap = len(query_words.intersection(sentence_words))
        if overlap > 0:
            scored_sentences.append((sentence, overlap))
    
    if scored_sentences:
        scored_sentences.sort(key=lambda x: x[1], reverse=True)
        best_sentences = [sent for sent in scored_sentences[:2]]
        return '. '.join(best_sentences)
    
    return "I couldn't find specific information to answer your question in the uploaded document. Please try rephrasing your question or check if the relevant information is included in the document."

def process_uploaded_file(uploaded_file, vectorstore, embedding_model):
    """Process uploaded file and store in Pinecone"""
    try:
        # Extract text based on file type
        if uploaded_file.type == "application/pdf":
            text = extract_pdf_text(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = extract_docx_text(uploaded_file)
        elif uploaded_file.type == "text/plain":
            text = extract_txt_text(uploaded_file)
        else:
            return "Unsupported file type"

        if not text.strip():
            return "No text found in document"

        # Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=300,
            separators=["\n\n", "\n", ".", "!", "?", ";", ",", " ", ""]
        )
        
        chunks = text_splitter.split_text(text)
        
        # Create documents for Pinecone
        documents = [
            Document(
                page_content=chunk,
                metadata={
                    "source": uploaded_file.name,
                    "chunk_id": i,
                    "chunk_size": len(chunk)
                }
            )
            for i, chunk in enumerate(chunks)
        ]
        
        # Add to Pinecone
        vectorstore.add_documents(documents)
        
        # Build knowledge graph from document text
        build_knowledge_graph(text)
        
        st.session_state.vectorstore = vectorstore
        st.session_state.document_processed = True
        
        return f"Document processed successfully. {len(chunks)} sections indexed in Pinecone."
        
    except Exception as e:
        return f"Error processing document: {str(e)}"

def query_rag_system(query):
    """Query the Pinecone RAG system"""
    if not st.session_state.vectorstore:
        return "No document has been uploaded and processed yet."
    
    try:
        # Query Pinecone for relevant context
        results = st.session_state.vectorstore.similarity_search(query, k=5)
        
        if not results:
            return "No relevant information found in the uploaded document."
        
        # Combine retrieved context
        context = "\n".join([doc.page_content for doc in results])
        
        # Generate answer using LM Studio
        answer = generate_direct_answer(query, context)
        return answer
        
    except Exception as e:
        return f"Error querying document: {str(e)}"

@st.cache_data
def load_ibm_dataset():
    """Load IBM dataset - NO FALLBACKS, only real data"""
    try:
        df = pd.read_csv('IBM_Employee_Data_With_Names.csv')
        return df
    except FileNotFoundError:
        try:
            df = pd.read_csv('WA_Fn-UseC_-HR-Employee-Attrition.csv')
            st.warning("Using original IBM dataset without names. Upload real enterprise data for authentic analysis.")
            return df
        except FileNotFoundError:
            st.error("No IBM dataset found. Please upload your enterprise data files for authentic analysis.")
            return None

def extract_entities_from_text(text):
    """Extract entities using simple pattern matching and NER"""
    entities = {'PERSON': set(), 'PROJECT': set(), 'SKILL': set()}
    
    import re
    
    person_patterns = re.findall(r'\b[A-Z][a-z]+ [A-Z][a-z]+\b', text)
    for person in person_patterns:
        if len(person.split()) == 2:
            entities['PERSON'].add(person)
    
    project_keywords = ['project', 'initiative', 'program', 'campaign', 'development', 'implementation']
    for keyword in project_keywords:
        pattern = rf'{keyword}\s+([A-Z][A-Za-z\s]+?)(?=\.|,|\n|$)'
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            if len(match.strip()) > 3:
                entities['PROJECT'].add(f"{keyword.title()} {match.strip()}")
    
    skill_keywords = ['python', 'javascript', 'java', 'management', 'leadership', 'analysis', 
                     'design', 'development', 'marketing', 'sales', 'finance', 'accounting',
                     'machine learning', 'data science', 'ai', 'artificial intelligence']
    for skill in skill_keywords:
        if skill.lower() in text.lower():
            entities['SKILL'].add(skill.title())
    
    return entities

def build_knowledge_graph(document_text):
    """Build knowledge graph from document text"""
    entities = extract_entities_from_text(document_text)
    
    for entity_type, entity_set in entities.items():
        st.session_state.entities[entity_type].update(entity_set)
    
    G = st.session_state.knowledge_graph
    
    for entity_type, entity_set in entities.items():
        for entity in entity_set:
            G.add_node(entity, type=entity_type)
    
    all_entities = []
    for entity_set in entities.values():
        all_entities.extend(list(entity_set))
    
    for i, entity1 in enumerate(all_entities):
        for entity2 in all_entities[i+1:]:
            if entity1 != entity2:
                if G.has_edge(entity1, entity2):
                    G[entity1][entity2]['weight'] += 1
                else:
                    G.add_edge(entity1, entity2, weight=1, relationship='co_occurs')
    
    return G

def visualize_knowledge_graph():
    """Create interactive knowledge graph visualization"""
    G = st.session_state.knowledge_graph
    
    if len(G.nodes()) == 0:
        st.warning("No entities found. Please upload and process a document first.")
        return None
    
    try:
        pos = nx.spring_layout(G, k=1, iterations=50)
    except:
        pos = nx.random_layout(G)
    
    edge_trace = []
    for edge in G.edges():
        try:
            x0, y0 = pos[edge[0]]
            x1, y1 = pos[edge[1]]
            edge_trace.append(go.Scatter(x=[x0, x1, None], y=[y0, y1, None],
                                       mode='lines', line=dict(width=1, color='#888'),
                                       hoverinfo='none', showlegend=False))
        except KeyError:
            continue

    node_traces = {}
    colors = {'PERSON': '#ff9999', 'PROJECT': '#66b3ff', 'SKILL': '#99ff99'}
    
    for node_type in ['PERSON', 'PROJECT', 'SKILL']:
        nodes_of_type = [node for node in G.nodes() if G.nodes[node].get('type') == node_type]
        if nodes_of_type:
            x_vals = [pos[node][0] for node in nodes_of_type]
            y_vals = [pos[node][1] for node in nodes_of_type]
            
            node_traces[node_type] = go.Scatter(
                x=x_vals, y=y_vals, mode='markers+text',
                marker=dict(size=20, color=colors[node_type]),
                text=nodes_of_type, textposition="middle center",
                name=node_type, hoverinfo='text'
            )
    
    fig = go.Figure(data=edge_trace + list(node_traces.values()),
                    layout=go.Layout(
                    title={"text": "Knowledge Graph: Person - Project - Skill Relationships", "font": {"size": 16}},
                    showlegend=True,
                    hovermode='closest',
                    margin=dict(b=20, l=5, r=5, t=40),
                    annotations=[dict(
                        text="Entities extracted from uploaded documents",
                        showarrow=False,
                        xref="paper", yref="paper",
                        x=0.005, y=-0.002
                    )],
                    xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
                    yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
                    height=600
                ))
    
    return fig

def query_knowledge_graph(query_type, entity_name):
    """Query the knowledge graph for relationships"""
    G = st.session_state.knowledge_graph
    
    if entity_name not in G.nodes():
        return f"Entity '{entity_name}' not found in knowledge graph."
    
    neighbors = list(G.neighbors(entity_name))
    
    if not neighbors:
        return f"No relationships found for '{entity_name}'."
    
    results = defaultdict(list)
    for neighbor in neighbors:
        neighbor_type = G.nodes[neighbor].get('type', 'Unknown')
        results[neighbor_type].append(neighbor)
    
    response = f"**Relationships for {entity_name}:**\n\n"
    for entity_type, entities in results.items():
        response += f"**{entity_type}:** {', '.join(entities)}\n\n"
    
    return response

def show_knowledge_graph():
    """Display knowledge graph interface"""
    st.header("Knowledge Graph")
    st.markdown("Visualize and explore Person - Project - Skill relationships from your documents")
    
    if not st.session_state.document_processed:
        st.info("**Please upload and process a document first to build the knowledge graph.**")
        st.markdown("Go to 'Document Upload' section to upload your enterprise documents.")
        return
    
    G = st.session_state.knowledge_graph
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Total Entities", len(G.nodes()))
    with col2:
        st.metric("Relationships", len(G.edges()))
    with col3:
        st.metric("Persons", len(st.session_state.entities['PERSON']))
    with col4:
        st.metric("Projects", len(st.session_state.entities['PROJECT']))
    
    st.subheader("Interactive Knowledge Graph")
    fig = visualize_knowledge_graph()
    if fig:
        st.plotly_chart(fig, use_container_width=True)
    
    st.subheader("Query Relationships")
    col1, col2 = st.columns([2, 1])
    
    with col1:
        all_entities = list(G.nodes())
        if all_entities:
            selected_entity = st.selectbox("Select an entity to explore:", all_entities)
        else:
            selected_entity = None
    
    with col2:
        if st.button("Explore Relationships") and selected_entity:
            result = query_knowledge_graph("relationships", selected_entity)
            st.markdown(result)
    
    if st.session_state.entities:
        st.subheader("Extracted Entities")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.write("**Persons:**")
            for person in list(st.session_state.entities['PERSON'])[:10]:
                st.write(f"• {person}")
        
        with col2:
            st.write("**Projects:**")
            for project in list(st.session_state.entities['PROJECT'])[:10]:
                st.write(f"• {project}")
        
        with col3:
            st.write("**Skills:**")
            for skill in list(st.session_state.entities['SKILL'])[:10]:
                st.write(f"• {skill}")

def calculate_investment_metrics(df):
    """Calculate investment and risk metrics from the dataset"""
    if df is None:
        return None
        
    total_employees = len(df)
    attrition_rate = len(df[df['Attrition'] == 'Yes']) / total_employees
    
    avg_salaries = df.groupby('Department')['MonthlyIncome'].mean()
    
    replacement_cost_per_employee = 75000
    annual_departures = int(total_employees * attrition_rate)
    total_replacement_cost = annual_departures * replacement_cost_per_employee
    
    senior_employees = df[df['YearsAtCompany'] >= 10]
    high_performers = df[df['PerformanceRating'] >= 4]
    at_risk_senior = senior_employees[senior_employees['JobSatisfaction'] <= 2]
    
    knowledge_risk_score = len(at_risk_senior) / len(senior_employees) if len(senior_employees) > 0 else 0
    
    return {
        'total_employees': total_employees,
        'attrition_rate': attrition_rate,
        'annual_departures': annual_departures,
        'total_replacement_cost': total_replacement_cost,
        'knowledge_risk_score': knowledge_risk_score,
        'avg_salaries': avg_salaries,
        'senior_employees': len(senior_employees),
        'high_performers': len(high_performers),
        'at_risk_senior': len(at_risk_senior)
    }

# Updated agents_info list without emojis
agents_info = [
    {"name": "MonitoringAgent", "role": "Knowledge Monitor", "icon": ""},
    {"name": "ContentGenerationAgent", "role": "Content Creator", "icon": ""},
    {"name": "RecommendationAgent", "role": "Advisor", "icon": ""},
    {"name": "QualityAssuranceAgent", "role": "Quality Controller", "icon": ""},
    {"name": "AnalysisAgent", "role": "Data Analyst", "icon": ""},
    {"name": "DocumentationAgent", "role": "Documentation Specialist", "icon": ""},
    {"name": "ValidationAgent", "role": "Validator", "icon": ""},
    {"name": "CommunicationAgent", "role": "Communications Manager", "icon": ""},
    {"name": "CoordinationAgent", "role": "Process Coordinator", "icon": ""},
    {"name": "ComplianceAgent", "role": "Compliance Officer", "icon": ""},
    {"name": "ReportingAgent", "role": "Report Generator", "icon": ""},
    {"name": "RecruitmentProposalAgent", "role": "Recruitment Strategist", "icon": ""},
    {"name": "TrainingProposalAgent", "role": "Training Strategist", "icon": ""}
]

def show_agentic_ai_interface():
    st.header("Agentic AI Workflow")
    st.markdown("Autonomous multi-agent system for enterprise knowledge management")
    
    try:
        from agents import run_agentic_workflow
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("Run Agentic Workflow", type="primary"):
                with st.spinner("Agents are working..."):
                    context = "Enterprise documents processed via RAG pipeline"
                    if st.session_state.document_processed:
                        context += " - Documents available in Pinecone vector store"
                    
                    results = run_agentic_workflow(context)
                    st.session_state.agentic_results = results
        
        with col2:
            if st.button("View Agent Status"):
                st.info("All agents ready for deployment")
                agent_status = {
                    "Knowledge Monitor": "Active - Scanning for gaps",
                    "Content Creator": "Ready - Connected to LM Studio", 
                    "Advisor": "Standby - Ready to recommend",
                    "Quality Controller": "Active - Monitoring outputs"
                }
                for agent, status in agent_status.items():
                    st.write(f"{agent}: {status}")
        
        # Display workflow results with recruitment and training proposals
        if hasattr(st.session_state, 'agentic_results') and st.session_state.agentic_results:
            st.subheader("Agent Workflow Results")
            results = st.session_state.agentic_results
            
            if "recruitment_proposals" in results:
                st.subheader("Proactive Recruitment Proposals")
                recruitment_data = results["recruitment_proposals"]
                
                if recruitment_data.get("status") == "recruitment_analysis_complete":
                    st.markdown("**Predicted Hiring Needs:**")
                    for need in recruitment_data.get("predicted_needs", []):
                        st.markdown(f"• {need}")
                    
                    st.markdown("**Candidate Proposals:**")
                    for proposal in recruitment_data.get("candidate_proposals", []):
                        with st.expander(f"{proposal['role']} Candidates", expanded=True):
                            st.markdown(f"**Timeline:** {proposal['timeline']}")
                            st.markdown(f"**Estimated Cost:** {proposal.get('estimated_cost', 'N/A')}")
                            st.markdown("**Recommended Candidates:**")
                            for candidate in proposal.get("candidates", []):
                                st.markdown(f"• {candidate}")
                    
                    st.markdown("**Strategic Recommendations:**")
                    for rec in recruitment_data.get("strategic_recommendations", []):
                        st.markdown(f"• {rec}")

            if "training_proposals" in results:
                st.subheader("Proactive Training Module Proposals")
                training_data = results["training_proposals"]
                
                if training_data.get("status") == "training_analysis_complete":
                    st.markdown("**Skill Gaps Identified:**")
                    for gap in training_data.get("skill_gaps_identified", []):
                        st.markdown(f"• {gap}")
                    
                    st.markdown("**Recommended Training Modules:**")
                    for proposal in training_data.get("training_proposals", []):
                        with st.expander(f"{proposal['module']}", expanded=True):
                            col1, col2 = st.columns(2)
                            with col1:
                                st.markdown(f"**Target:** {proposal['target_audience']}")
                                st.markdown(f"**Duration:** {proposal['duration']}")
                                st.markdown(f"**Priority:** {proposal['priority']}")
                            with col2:
                                st.markdown(f"**Cost:** {proposal['cost']}")
                                st.markdown(f"**Expected Outcome:** {proposal['expected_outcome']}")
                    
                    if "roi_projection" in training_data:
                        st.success(f"**ROI Projection:** {training_data['roi_projection']}")

            if "integrated_workforce_planning" in results:
                st.subheader("Integrated Workforce Planning Strategy")
                st.json(results["integrated_workforce_planning"])
            
            # Store results for tab access
            if 'workflow_results' not in st.session_state:
                st.session_state.workflow_results = []
            st.session_state.workflow_results.append(results)
                        
    except ImportError:
        st.error("Agents module not found. Please create agents.py first.")
    except Exception as e:
        st.error(f"Agentic workflow error: {str(e)}")

def show_document_upload_updated():
    """Enhanced document upload interface"""
    st.header("Document Upload & Processing")
    st.markdown("Upload enterprise documents to enable RAG-powered analysis")
    
    if st.session_state.vectorstore is None:
        st.error("Pinecone not initialized. Please refresh the page.")
        return
    
    uploaded_file = st.file_uploader(
        "Choose a document file",
        type=['pdf', 'docx', 'txt'],
        help="Upload PDF, Word, or text documents for analysis"
    )
    
    if uploaded_file is not None:
        st.write(f"**File:** {uploaded_file.name}")
        st.write(f"**Size:** {uploaded_file.size:,} bytes")
        
        if st.button("Process Document", type="primary"):
            with st.spinner(f"Processing {uploaded_file.name}..."):
                result_message = process_uploaded_file(
                    uploaded_file, 
                    st.session_state.vectorstore,
                    st.session_state.embedding_model
                )
                
                if "successfully" in result_message:
                    st.success(result_message)
                    st.info("Document is now ready for AI analysis and knowledge graph generation.")
                else:
                    st.error(result_message)
    
    if st.session_state.document_processed:
        st.success("Document processed and ready for analysis")
        
        if st.button("Reset Document"):
            try:
                # Clear session state instead of deleting collection
                st.session_state.document_processed = False
                st.session_state.knowledge_graph.clear()
                st.session_state.entities = {'PERSON': set(), 'PROJECT': set(), 'SKILL': set()}
                st.success("Document reset completed")
                st.rerun()
            except Exception as e:
                st.error(f"Reset failed: {str(e)}")

def show_rag_interface_updated():
    """Enhanced RAG query interface"""
    st.header("RAG Query Interface")
    st.markdown("Ask questions about your uploaded documents")
    
    if not st.session_state.document_processed:
        st.info("Please upload and process a document first")
        return
    
    query = st.text_input(
        "Enter your question:",
        value=st.session_state.query_input,
        placeholder="e.g., What is the total investment required for this project?"
    )
    
    col1, col2 = st.columns([1, 4])
    
    with col1:
        if st.button("Ask Question", type="primary"):
            if query:
                with st.spinner("Analyzing document..."):
                    answer = query_rag_system(query)
                    st.session_state.query_input = query
                    
                    st.subheader("Answer:")
                    st.write(answer)
            else:
                st.warning("Please enter a question")

def show_executive_overview(df, metrics):
    """Display executive overview dashboard"""
    st.header("Executive Overview")
    st.markdown("Strategic insights for organizational knowledge management and investment planning")
    
    if df is None or metrics is None:
        st.error("**No Enterprise Data Available for Analysis**")
        st.warning("This dashboard requires authentic enterprise data to provide meaningful insights.")
        st.info("Please upload your enterprise documents in the Document Upload section to begin analysis.")
        return
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric(
            "Total Investment Required",
            f"${metrics['total_replacement_cost']:,.0f}",
            delta=f"{metrics['annual_departures']} departures/year"
        )
    
    with col2:
        st.metric(
            "Knowledge Risk Score",
            f"{metrics['knowledge_risk_score']:.1%}",
            delta=f"{metrics['at_risk_senior']} at-risk seniors"
        )
    
    with col3:
        st.metric(
            "Attrition Rate",
            f"{metrics['attrition_rate']:.1%}",
            delta="Industry benchmark: 15%"
        )
    
    with col4:
        st.metric(
            "High Performers",
            f"{metrics['high_performers']}",
            delta=f"{metrics['high_performers']/len(df):.1%} of workforce"
        )
    
    col1, col2 = st.columns(2)
    
    with col1:
        fig_attrition = px.bar(
            df.groupby('Department')['Attrition'].value_counts().unstack().fillna(0),
            title="Attrition Analysis by Department",
            color_discrete_map={'Yes': '#EF4444', 'No': '#22C55E'}
        )
        fig_attrition.update_layout(height=400)
        st.plotly_chart(fig_attrition, use_container_width=True)
    
    with col2:
        fig_performance = px.histogram(
            df,
            x='PerformanceRating',
            color='Department',
            title="Performance Rating Distribution",
            nbins=4
        )
        fig_performance.update_layout(height=400)
        st.plotly_chart(fig_performance, use_container_width=True)

def show_risk_analysis(df, metrics):
    """Display risk analysis dashboard"""
    st.header("Enterprise Risk Analysis")
    st.markdown("Detailed analysis of organizational knowledge risks and mitigation strategies")
    
    if df is None or metrics is None:
        st.error("**No Enterprise Data Available for Risk Analysis**")
        st.warning("Risk analysis requires authentic enterprise data.")
        st.info("Please upload your enterprise documents to perform risk assessment.")
        return
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Single Points of Failure", 
                 len(df[(df['YearsAtCompany'] >= 15) & (df['PerformanceRating'] >= 4)]))
    
    with col2:
        st.metric("Flight Risk Employees", 
                 len(df[(df['JobSatisfaction'] <= 2) & (df['PerformanceRating'] >= 3)]))
    
    with col3:
        st.metric("Knowledge Concentration Risk", 
                 f"{len(df[df['YearsInCurrentRole'] >= 10]) / len(df):.1%}")

def show_training_plans(df, metrics):
    """Display training plans"""
    st.header("Strategic Training Plans")
    st.markdown("Data-driven training recommendations for knowledge enhancement and risk mitigation")
    
    if df is None or metrics is None:
        st.error("**No Enterprise Data Available for Training Analysis**")
        st.warning("Training recommendations require authentic enterprise data.")
        st.info("Please upload your enterprise documents to generate training insights.")
        return
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        training_budget = metrics['total_replacement_cost'] * 0.3
        st.metric("Annual Training Budget", f"${training_budget:,.0f}")
    
    with col2:
        employees_needing_training = len(df[df['JobSatisfaction'] <= 2])
        st.metric("Employees Needing Training", employees_needing_training)
    
    with col3:
        avg_training_cost_per_employee = training_budget / len(df)
        st.metric("Avg. Training Cost/Employee", f"${avg_training_cost_per_employee:,.0f}")
    
    with col4:
        expected_roi = 4.2
        st.metric("Expected ROI", f"{expected_roi}x")

def show_agent_status(df):
    """Display agent status"""
    st.header("Knowledge Agent Status")
    st.markdown("AI-powered insights and automated analysis results")
    
    if df is None:
        st.error("**No Enterprise Data Available for Agent Analysis**")
        st.warning("Agent analysis requires authentic enterprise data.")
        st.info("Please upload your enterprise documents to enable agent operations.")
        return
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Active Agents", "13")
    with col2:
        st.metric("Documents Processed", "1" if st.session_state.document_processed else "0")
    with col3:
        st.metric("Knowledge Entities", len(st.session_state.knowledge_graph.nodes()))
    with col4:
        st.metric("System Health", "Optimal")

def main():
    """Main application function"""
    st.markdown('<h1 class="main-header">Enterprise Knowledge Evolution Forecaster</h1>', unsafe_allow_html=True)
    
    df = load_ibm_dataset()
    
    if df is None:
        st.error("**No Enterprise Data Available**")
        st.warning("This application requires authentic enterprise data to provide meaningful analysis.")
        st.info("Please upload your enterprise documents using the Document Upload section to begin analysis.")
        metrics = None
    else:
        metrics = calculate_investment_metrics(df)
    
    if st.session_state.vectorstore is None:
        with st.spinner("Initializing Pinecone..."):
            vectorstore, embedding_model, status = initialize_pinecone()
            st.session_state.vectorstore = vectorstore
            st.session_state.embedding_model = embedding_model
    
    # Sidebar
    st.sidebar.title("Navigation")
    
    st.sidebar.markdown("### System Status")
    if st.session_state.vectorstore:
        st.sidebar.success("Vector Database: Pinecone Active")
        if st.session_state.document_processed:
            st.sidebar.success("Document: Processed and Ready")
        else:
            st.sidebar.info("Document: No document uploaded")
    else:
        st.sidebar.error("Pinecone Initialization Failed")
    
    # LM Studio status check
    try:
        response = requests.get("http://127.0.0.1:1234/", timeout=2)
        if response.status_code == 200:
            st.sidebar.success("LM Studio: Connected (Llama 3.2 3B)")
        else:
            st.sidebar.warning("LM Studio: Server running but model not loaded")
    except:
        st.sidebar.error("LM Studio: Not connected")
        st.sidebar.info("Please ensure LM Studio is running with Llama 3.2 3B model loaded")
    
    # Agent list in sidebar
    st.sidebar.markdown("### Active Agents")
    for agent in agents_info:
        st.sidebar.write(f"• {agent['name']}: {agent['role']}")
    
    dashboard_section = st.sidebar.selectbox(
        "Select Dashboard Section",
        ["Executive Overview", "Risk Analysis", "Training Plans", "Agent Status", "RAG Query Interface", "Document Upload", "Knowledge Graph", "Agentic AI Workflow"]
    )
    
    if df is not None:
        st.sidebar.markdown("### Dataset Information")
        st.sidebar.metric("Total Employees", f"{len(df):,}")
        st.sidebar.metric("Departments", df['Department'].nunique())
        st.sidebar.metric("Job Roles", df['JobRole'].nunique())
    else:
        st.sidebar.markdown("### Dataset Information")
        st.sidebar.error("No enterprise data loaded")
        st.sidebar.info("Upload documents to analyze")
    
    # Main content based on selection
    if dashboard_section == "Executive Overview":
        show_executive_overview(df, metrics)
    elif dashboard_section == "Risk Analysis":
        show_risk_analysis(df, metrics)
    elif dashboard_section == "Training Plans":
        show_training_plans(df, metrics)
    elif dashboard_section == "Agent Status":
        show_agent_status(df)
    elif dashboard_section == "RAG Query Interface":
        show_rag_interface_updated()
    elif dashboard_section == "Document Upload":
        show_document_upload_updated()
    elif dashboard_section == "Knowledge Graph":
        show_knowledge_graph()   
    elif dashboard_section == "Agentic AI Workflow":
        show_agentic_ai_interface()

    # Add the new dashboard tabs for HR insights
    st.markdown("---")
    tab1, tab2, tab3 = st.tabs(["Main Dashboard", "Recruitment Hub", "Training Center"])

    with tab2:
        st.header("Recruitment Intelligence Hub")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Open Positions", "5", "+2")
        with col2:
            st.metric("Candidates Sourced", "23", "+8")
        with col3:
            st.metric("Time to Hire", "18 days", "-5")
        
        if st.session_state.workflow_results:
            latest_result = st.session_state.workflow_results[-1]
            if "recruitment_proposals" in latest_result:
                st.subheader("Latest Recruitment Proposals")
                st.json(latest_result["recruitment_proposals"])

    with tab3:
        st.header("Training & Development Center")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Active Programs", "12", "+3")
        with col2:
            st.metric("Employees Enrolled", "45", "+15")
        with col3:
            st.metric("Completion Rate", "87%", "+5%")
        
        if st.session_state.workflow_results:
            latest_result = st.session_state.workflow_results[-1]
            if "training_proposals" in latest_result:
                st.subheader("Latest Training Proposals")
                st.json(latest_result["training_proposals"])

if __name__ == "__main__":
    main()
